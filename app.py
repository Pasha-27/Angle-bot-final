import streamlit as st
import requests
import pandas as pd
import re
import json
import time
import base64
import os
import subprocess
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Tuple, Union
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import docx2txt
from pydub import AudioSegment  # New import for audio chunking
import io

# Configure page settings
st.set_page_config(
    page_title="YouTube Downloader",
    page_icon="💬",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={
        'About': "# YouTube Downloader\nDownload comments and transcripts from multiple YouTube videos."
    }
)

# Initialize session state for persistence across reruns
if 'processed_videos' not in st.session_state:
    st.session_state.processed_videos = []

if 'analyzing_video' not in st.session_state:
    st.session_state.analyzing_video = None

if 'analysis_results' not in st.session_state:
    st.session_state.analysis_results = {}

# Load API keys from st.secrets if available
@st.cache_resource
def get_api_keys():
    config = {"youtube_api_key": None, "openai_api_key": None}
    try:
        if "YOUTUBE_API_KEY" in st.secrets:
            config["youtube_api_key"] = st.secrets["YOUTUBE_API_KEY"]
        if "OPENAI_API_KEY" in st.secrets:
            config["openai_api_key"] = st.secrets["OPENAI_API_KEY"]
    except Exception:
        pass
    return config

api_keys = get_api_keys()
YOUTUBE_API_KEY = api_keys["youtube_api_key"]
OPENAI_API_KEY = api_keys["openai_api_key"]

# Create downloads directory if it doesn't exist
os.makedirs("downloads", exist_ok=True)

# Load angle_bot_prompt.txt
def load_analysis_prompt():
    try:
        with open("angle_bot_prompt.txt", "r", encoding="utf-8") as f:
            return f.read()
    except Exception as e:
        st.warning(f"Could not load angle_bot_prompt.txt: {str(e)}. Using default prompt.")
        return "Analyze the transcript and comments from this YouTube video. Identify key themes, interesting insights, and what seems to resonate with viewers."

ANALYSIS_PROMPT = load_analysis_prompt()

# Apply custom CSS for dark mode and modern UI
st.markdown("""
<style>
    /* Dark mode theme */
    .stApp {
        background-color: #121212;
        color: #f0f0f0;
    }
    /* Card-like containers */
    .card {
        background-color: #1e1e1e;
        border-radius: 10px;
        padding: 20px;
        margin-bottom: 20px;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    /* Header styling */
    h1, h2, h3 {
        color: #ffffff;
        font-weight: 600;
    }
    /* Accent colors for highlights */
    .highlight {
        color: #ff5252;
    }
    /* Button hover effects */
    .stButton>button:hover {
        background-color: #ff5252;
        border-color: #ff5252;
    }
    /* Progress bar styling */
    .stProgress > div > div {
        background-color: #ff5252;
    }
    /* Custom divider */
    .divider {
        height: 3px;
        background: linear-gradient(90deg, #ff5252, transparent);
        margin: 20px 0;
        border-radius: 10px;
    }
    /* Download button styling */
    .download-button {
        display: inline-block;
        background-color: #ff5252;
        color: white;
        padding: 8px 16px;
        border-radius: 4px;
        text-decoration: none;
        font-weight: bold;
        margin: 10px 0;
        transition: background-color 0.3s;
    }
    .download-button:hover {
        background-color: #ff3333;
    }
</style>
""", unsafe_allow_html=True)

# Helper function to remove invalid filename characters
def clean_filename(filename: str) -> str:
    cleaned = re.sub(r'[\\/*?:"<>|]', "_", filename)
    return cleaned.strip('. ')[:100]

# Helper function: extract video ID from a YouTube URL
def extract_video_id(url: str) -> Optional[str]:
    patterns = [
        r"(?:v=|\/)([0-9A-Za-z_-]{11}).*",
        r"(?:youtu\.be\/)([0-9A-Za-z_-]{11})",
        r"(?:embed\/)([0-9A-Za-z_-]{11})"
    ]
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    return None

# Helper function: get binary file download link
def get_binary_file_downloader_html(bin_file, file_label='File'):
    """Generate a link to download a binary file."""
    with open(bin_file, 'rb') as f:
        data = f.read()
    bin_str = base64.b64encode(data).decode()
    return f'<a href="data:application/octet-stream;base64,{bin_str}" download="{os.path.basename(bin_file)}" class="download-button">{file_label}</a>'

# Helper function: fetch comments for a video using the YouTube Data API v3
def get_comments(video_id: str, api_key: str, page_token: Optional[str] = None) -> Dict:
    api_url = "https://www.googleapis.com/youtube/v3/commentThreads"
    params = {
        "part": "snippet,replies",
        "videoId": video_id,
        "maxResults": 100,
        "textFormat": "plainText",
        "key": api_key,
        "order": "relevance"
    }
    if page_token:
        params["pageToken"] = page_token
    try:
        response = requests.get(api_url, params=params, timeout=10)
        if response.status_code != 200:
            error_message = f"HTTP error {response.status_code}"
            try:
                error_data = response.json()
                if "error" in error_data and "message" in error_data["error"]:
                    error_message = error_data["error"]["message"]
            except Exception:
                pass
            return {"success": False, "error": error_message}
        return {"success": True, "data": response.json()}
    except requests.exceptions.Timeout:
        return {"success": False, "error": "Request timed out. YouTube API might be slow or experiencing issues."}
    except requests.exceptions.ConnectionError:
        return {"success": False, "error": "Connection error. Please check your internet connection."}
    except Exception as e:
        return {"success": False, "error": f"Unexpected error: {str(e)}"}

# Helper function: process comment data from the API response
def process_comment_data(data: Dict) -> Tuple[List[Dict], Optional[str]]:
    comments = []
    for item in data.get("items", []):
        snippet = item["snippet"]["topLevelComment"]["snippet"]
        comments.append({
            "author": snippet.get("authorDisplayName", "Unknown"),
            "text": snippet.get("textDisplay", ""),
            "likeCount": snippet.get("likeCount", 0),
            "publishedAt": snippet.get("publishedAt", ""),
            "isReply": False
        })
        if "replies" in item:
            for reply in item["replies"]["comments"]:
                reply_snippet = reply["snippet"]
                comments.append({
                    "author": reply_snippet.get("authorDisplayName", "Unknown"),
                    "text": reply_snippet.get("textDisplay", ""),
                    "likeCount": reply_snippet.get("likeCount", 0),
                    "publishedAt": reply_snippet.get("publishedAt", ""),
                    "isReply": True
                })
    next_page_token = data.get("nextPageToken")
    return comments, next_page_token

# Helper function: format API timestamps to a readable date
def format_timestamp(timestamp: str) -> str:
    try:
        dt = datetime.strptime(timestamp, "%Y-%m-%dT%H:%M:%SZ")
        return dt.strftime("%b %d, %Y")
    except Exception:
        return timestamp

# Helper function: get basic video information
def get_video_info(video_id: str, api_key: str) -> Dict:
    api_url = "https://www.googleapis.com/youtube/v3/videos"
    params = {
        "part": "snippet,statistics",
        "id": video_id,
        "key": api_key
    }
    response = requests.get(api_url, params=params)
    if response.status_code != 200:
        return {"success": False, "error": f"HTTP error {response.status_code}: {response.text}"}
    data = response.json()
    if not data.get("items"):
        return {"success": False, "error": "Video not found"}
    video_data = data["items"][0]
    snippet = video_data["snippet"]
    statistics = video_data["statistics"]
    return {
        "success": True,
        "title": snippet.get("title", "Unknown Title"),
        "channel": snippet.get("channelTitle", "Unknown Channel"),
        "published": snippet.get("publishedAt", "Unknown Date"),
        "comment_count": statistics.get("commentCount", "0"),
        "view_count": statistics.get("viewCount", "0"),
        "like_count": statistics.get("likeCount", "0"),
        "thumbnail": snippet.get("thumbnails", {}).get("high", {}).get("url", "")
    }

# Helper function: fetch all comments with pagination
def get_all_comments_with_callback(video_id: str, api_key: str, status_text=None) -> List[Dict]:
    all_comments = []
    next_page_token = None
    page_num = 1
    max_pages = 500
    error_count = 0
    max_errors = 3
    while page_num <= max_pages:
        if status_text:
            status_text.text(f"Fetching page {page_num} of comments...")
        result = get_comments(video_id, api_key, next_page_token)
        if not result["success"]:
            error_count += 1
            if error_count >= max_errors:
                if status_text:
                    status_text.error("Too many errors. Stopping comment retrieval.")
                break
            time.sleep(2)
            continue
        error_count = 0
        comments, next_page_token = process_comment_data(result["data"])
        all_comments.extend(comments)
        if not next_page_token:
            break
        time.sleep(0.5)
        page_num += 1
    return all_comments

# Function to download audio from a YouTube video
def download_youtube_audio(youtube_url, output_path="./downloads", format="mp3", quality="192"):
    """Download audio from a YouTube video using yt-dlp."""
    try:
        video_id = extract_video_id(youtube_url)
        if not video_id:
            return {"success": False, "error": "Could not extract video ID from URL"}
            
        video_info = get_video_info(video_id, YOUTUBE_API_KEY)
        if not video_info["success"]:
            return video_info
            
        title = video_info["title"]
        base_filename = clean_filename(title)
        output_file = f"{output_path}/{base_filename}.{format}"
        
        command = [
            "yt-dlp", "-f", "bestaudio",
            "--extract-audio", "--audio-format", format,
            "--audio-quality", quality,
            "-o", output_file,
            "--no-playlist", "--quiet", "--progress",
            youtube_url
        ]
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        status_text.text("Downloading audio...")
        
        process = subprocess.run(command, capture_output=True, text=True)
        
        if process.returncode != 0:
            return {"success": False, "error": process.stderr}
            
        progress_bar.progress(100)
        status_text.empty()
        
        if not os.path.exists(output_file):
            potential_files = list(Path(output_path).glob(f"{base_filename}.*"))
            if potential_files:
                output_file = str(potential_files[0])
            else:
                return {"success": False, "error": "File not found after download"}
                
        return {
            "success": True,
            "file_path": output_file,
            "title": title,
            "file_name": os.path.basename(output_file)
        }
    except Exception as e:
        return {"success": False, "error": str(e)}

# Function to transcribe audio using OpenAI's Whisper
def transcribe_with_whisper(audio_file_path, api_key):
    """Convert audio to text using OpenAI's Whisper API.
    If the audio file size exceeds 25MB, break it into smaller chunks (up to 5 minutes per chunk)
    and transcribe each chunk sequentially."""
    try:
        import openai
        openai.api_key = api_key

        size_limit = 25 * 1024 * 1024  # 25 MB in bytes
        file_size = os.path.getsize(audio_file_path)
        
        # If file size is within the limit, transcribe normally.
        if file_size <= size_limit:
            with open(audio_file_path, "rb") as audio_file:
                transcript_response = openai.Audio.transcribe("whisper-1", audio_file)
            transcript_text = transcript_response.get("text", "")
            return {"success": True, "transcript": transcript_text}
        else:
            st.info("Audio file is larger than 25MB. Splitting into smaller chunks for transcription...")
            # Determine file extension for pydub
            ext = os.path.splitext(audio_file_path)[1][1:]
            # Load audio with pydub
            audio = AudioSegment.from_file(audio_file_path, format=ext)
            duration_ms = len(audio)
            # Calculate average bytes per ms from the original file
            avg_bytes_per_ms = file_size / duration_ms
            # Maximum chunk duration in ms so that the chunk size is under the 25MB limit
            max_chunk_duration_ms = int(size_limit / avg_bytes_per_ms)
            # To optimize API usage, cap chunk duration at 5 minutes (300000 ms)
            chunk_duration_ms = min(max_chunk_duration_ms, 5 * 60 * 1000)
            
            transcript_text = ""
            num_chunks = (duration_ms // chunk_duration_ms) + (1 if duration_ms % chunk_duration_ms != 0 else 0)
            st.write(f"Processing {num_chunks} chunk(s) ...")
            for i in range(0, duration_ms, chunk_duration_ms):
                chunk = audio[i: i + chunk_duration_ms]
                buf = io.BytesIO()
                # Export chunk to a BytesIO object in mp3 format
                chunk.export(buf, format="mp3")
                buf.seek(0)
                
                st.write(f"Transcribing chunk {(i // chunk_duration_ms) + 1} of {num_chunks} ...")
                result = openai.Audio.transcribe("whisper-1", buf)
                chunk_text = result.get("text", "")
                transcript_text += chunk_text + " "
            return {"success": True, "transcript": transcript_text}
    except Exception as e:
        return {"success": False, "error": f"Error transcribing audio with OpenAI: {str(e)}"}
        
# Function to extract text content from DOCX file
def extract_text_from_docx(docx_path):
    """Extract all text content from a DOCX file."""
    try:
        return docx2txt.process(docx_path)
    except Exception as e:
        st.error(f"Error extracting text from document: {str(e)}")
        return None
        
# Function to analyze document with OpenAI's ChatGPT
def analyze_document_with_chatgpt(document_text, prompt_text, api_key):
    """Analyze document content using ChatGPT with text chunking for long content."""
    try:
        import openai
        openai.api_key = api_key
        
        # Combine the prompt and truncated document text
        total_text = document_text
        
        # If document is too long, we need to chunk it
        # First, try GPT-4 as it has a larger context window
        try:
            # Try with GPT-4 (larger context window)
            full_prompt = f"{prompt_text}\n\nHere is the document containing the transcript and comments:\n\n{total_text}"
            
            response = openai.ChatCompletion.create(
                model="gpt-4",  # Using GPT-4 for larger context
                messages=[
                    {"role": "system", "content": "You are an expert analyst of YouTube video content and audience engagement."},
                    {"role": "user", "content": full_prompt}
                ],
                max_tokens=4000,
                temperature=0.7
            )
            
            analysis = response.choices[0].message.content
            return {"success": True, "analysis": analysis}
            
        except Exception as e:
            # If GPT-4 fails or is not available, try chunking with GPT-3.5
            if "maximum context length" in str(e) or "not found" in str(e):
                # Split the document into chunks
                # Extract transcript and comments sections
                if "## Transcript" in total_text and "## Comments" in total_text:
                    # Split by sections
                    parts = total_text.split("## Transcript", 1)
                    header = parts[0]
                    rest = "## Transcript" + parts[1]
                    
                    transcript_comments_parts = rest.split("## Comments", 1)
                    transcript_section = transcript_comments_parts[0]
                    comments_section = "## Comments" + transcript_comments_parts[1]
                    
                    # First, analyze the transcript
                    transcript_prompt = f"{prompt_text}\n\nHere is the TRANSCRIPT ONLY from the YouTube video:\n\n{header}\n{transcript_section}\n\nPlease analyze the transcript content. I will provide the comments separately."
                    
                    transcript_response = openai.ChatCompletion.create(
                        model="gpt-3.5-turbo-16k",
                        messages=[
                            {"role": "system", "content": "You are an expert analyst of YouTube video content."},
                            {"role": "user", "content": transcript_prompt}
                        ],
                        max_tokens=2000,
                        temperature=0.7
                    )
                    
                    transcript_analysis = transcript_response.choices[0].message.content
                    
                    # Now, analyze a sample of comments
                    # Extract some representative comments (first few, some middle, some end)
                    comment_lines = comments_section.split('\n')
                    # Take about 200 lines max spread throughout the comments
                    if len(comment_lines) > 200:
                        sample_comments = '\n'.join(
                            comment_lines[:50] +  # First 50 lines
                            comment_lines[len(comment_lines)//2-25:len(comment_lines)//2+25] +  # Middle 50 lines
                            comment_lines[-50:]  # Last 50 lines
                        )
                    else:
                        sample_comments = comments_section
                    
                    comments_prompt = f"{prompt_text}\n\nHere is a SAMPLE OF COMMENTS from the YouTube video (transcript already analyzed separately):\n\n{sample_comments}\n\nPlease analyze these comments and identify patterns, reactions, and viewer sentiment."
                    
                    comments_response = openai.ChatCompletion.create(
                        model="gpt-3.5-turbo-16k",
                        messages=[
                            {"role": "system", "content": "You are an expert analyst of audience engagement and comments."},
                            {"role": "user", "content": comments_prompt}
                        ],
                        max_tokens=2000,
                        temperature=0.7
                    )
                    
                    comments_analysis = comments_response.choices[0].message.content
                    
                    # Final synthesis
                    synthesis_prompt = f"Based on the following separate analyses of a YouTube video's transcript and comments, provide a comprehensive analysis that integrates both perspectives:\n\n1. TRANSCRIPT ANALYSIS:\n{transcript_analysis}\n\n2. COMMENTS ANALYSIS:\n{comments_analysis}\n\nPlease synthesize these into a cohesive overall analysis."
                    
                    synthesis_response = openai.ChatCompletion.create(
                        model="gpt-3.5-turbo-16k",
                        messages=[
                            {"role": "system", "content": "You are an expert analyst of YouTube video content and audience engagement."},
                            {"role": "user", "content": synthesis_prompt}
                        ],
                        max_tokens=2000,
                        temperature=0.7
                    )
                    
                    final_analysis = synthesis_response.choices[0].message.content
                    return {"success": True, "analysis": final_analysis}
                else:
                    # Fallback if we can't identify sections clearly
                    # Just truncate the document to fit
                    # Rough estimate: 1 token ≈ 4 chars
                    max_chars = 14000 * 4  # Leave room for prompt
                    truncated_text = total_text[:max_chars] + "\n...[Content truncated due to length]..."
                    
                    truncated_prompt = f"{prompt_text}\n\nHere is a truncated version of the document (it was too long for complete analysis):\n\n{truncated_text}"
                    
                    truncated_response = openai.ChatCompletion.create(
                        model="gpt-3.5-turbo-16k",
                        messages=[
                            {"role": "system", "content": "You are an expert analyst of YouTube video content and audience engagement."},
                            {"role": "user", "content": truncated_prompt}
                        ],
                        max_tokens=2000,
                        temperature=0.7
                    )
                    
                    truncated_analysis = truncated_response.choices[0].message.content
                    return {"success": True, "analysis": truncated_analysis + "\n\n[Note: This analysis is based on a truncated version of the document as the full content exceeded the model's context limits.]"}
            else:
                # If it's a different error, raise it
                raise e
                
    except Exception as e:
        return {"success": False, "error": f"Error analyzing document with ChatGPT: {str(e)}"}

# Function to create a Word document with transcript and comments
def create_word_doc(video_info, transcript, comments, output_path="./downloads"):
    """Create a Word document with the video transcript and comments."""
    try:
        # Create a new Document
        doc = Document()
        
        # Add document title
        title = doc.add_heading(f"{video_info['title']}", level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add video metadata
        doc.add_paragraph(f"Channel: {video_info['channel']}")
        doc.add_paragraph(f"Published: {video_info['published']}")
        
        # Add transcript section
        doc.add_heading("Transcript", level=2)
        doc.add_paragraph(transcript)
        
        # Add comments section
        doc.add_heading(f"Comments ({len(comments)})", level=2)
        
        # Add table for comments
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Set table header
        header_cells = table.rows[0].cells
        header_cells[0].text = "Author"
        header_cells[1].text = "Comment"
        header_cells[2].text = "Likes"
        
        # Add comments to the table
        for comment in comments:
            row_cells = table.add_row().cells
            row_cells[0].text = comment["author"]
            row_cells[1].text = comment["text"]
            row_cells[2].text = str(comment["likeCount"])
        
        # Save the document
        safe_title = clean_filename(video_info["title"])
        file_path = f"{output_path}/{safe_title}_transcript_comments.docx"
        doc.save(file_path)
        
        return {"success": True, "file_path": file_path}
    except Exception as e:
        return {"success": False, "error": f"Error creating Word document: {str(e)}"}

# Function to handle the analysis of a document
def handle_analyze_document(video_id, doc_path, safe_title):
    try:
        # Extract text from document
        doc_text = extract_text_from_docx(doc_path)
        
        if not doc_text:
            return {"success": False, "error": "Could not extract text from document"}
        
        # Analyze with ChatGPT
        analysis_result = analyze_document_with_chatgpt(
            doc_text, 
            ANALYSIS_PROMPT, 
            OPENAI_API_KEY or st.session_state.get("openai_key", "")
        )
        
        if analysis_result["success"]:
            # Save analysis to file
            analysis_text = analysis_result["analysis"]
            analysis_filename = f"{safe_title}_analysis.txt"
            analysis_path = f"downloads/{analysis_filename}"
            
            with open(analysis_path, "w", encoding="utf-8") as f:
                f.write(analysis_text)
                
            return {
                "success": True, 
                "analysis": analysis_result["analysis"],
                "analysis_path": analysis_path,
                "analysis_filename": analysis_filename
            }
        else:
            return analysis_result
    except Exception as e:
        return {"success": False, "error": f"Error during analysis: {str(e)}"}

# Main Application UI
st.markdown('<h1 class="highlight">YouTube Comments & Transcript Downloader</h1>', unsafe_allow_html=True)
st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

# API Key Input Section
with st.expander("API Keys (Required)", expanded=not (YOUTUBE_API_KEY and OPENAI_API_KEY)):
    col1, col2 = st.columns(2)
    
    with col1:
        youtube_api_key_input = st.text_input(
            "YouTube API Key", 
            value=YOUTUBE_API_KEY if YOUTUBE_API_KEY else "",
            type="password", 
            help="Required for fetching video info and comments",
            key="youtube_key"
        )
    
    with col2:
        openai_api_key_input = st.text_input(
            "OpenAI API Key", 
            value=OPENAI_API_KEY if OPENAI_API_KEY else "",
            type="password", 
            help="Required for audio transcription with Whisper and analysis with ChatGPT",
            key="openai_key"
        )
    
    if not youtube_api_key_input:
        st.warning("⚠️ YouTube API key is required.")
    
    if not openai_api_key_input:
        st.warning("⚠️ OpenAI API key is required for transcription and analysis.")

# Create tabs for different sections of the app
tab1, tab2 = st.tabs(["Process Videos", "Analyze Videos"])

with tab1:
    # Input area for multiple YouTube video URLs
    st.markdown("### Enter up to 5 YouTube video URLs (one per line):")
    video_urls_text = st.text_area(
        "YouTube Video URLs", 
        placeholder="https://www.youtube.com/watch?v=...\nhttps://www.youtube.com/watch?v=...",
        key="video_urls"
    )
    
    # Process button
    if st.button("Process Videos", key="process_btn"):
        # Validate API keys
        if not youtube_api_key_input:
            st.error("YouTube API key is required. Please enter your API key in the API Keys section.")
            st.stop()
        
        if not openai_api_key_input:
            st.error("OpenAI API key is required. Please enter your API key in the API Keys section.")
            st.stop()
        
        # Split the input text into lines and take up to 5 non-empty URLs
        video_urls = [url.strip() for url in video_urls_text.splitlines() if url.strip()]
        
        if not video_urls:
            st.error("Please enter at least one YouTube URL.")
        elif len(video_urls) > 5:
            st.error("Please enter no more than 5 URLs. Processing multiple videos consumes significant API resources.")
        else:
            # Clear previously processed videos
            st.session_state.processed_videos = []
            
            for url in video_urls:
                with st.spinner(f"Processing {url} ..."):
                    video_id = extract_video_id(url)
                    if not video_id:
                        st.error(f"Could not extract video ID from the URL: {url}")
                        continue
                    
                    # Step 1: Fetch video info
                    video_info = get_video_info(video_id, youtube_api_key_input)
                    if not video_info.get("success"):
                        st.error(f"Error fetching video info: {video_info.get('error', 'Unknown error')}")
                        continue
                    
                    # Display video info
                    st.subheader(f"📺 {video_info['title']}")
                    st.write(f"Channel: {video_info['channel']}")
                    st.write(f"View Count: {int(video_info['view_count']):,}")
                    
                    # Step 2: Download audio
                    audio_status = st.empty()
                    audio_status.info("Downloading audio...")
                    audio_result = download_youtube_audio(url)
                    
                    if not audio_result.get("success"):
                        st.error(f"Error downloading audio: {audio_result.get('error', 'Unknown error')}")
                        continue
                    
                    audio_status.success("✅ Audio downloaded successfully")
                    
                    # Step 3: Generate transcript with Whisper
                    transcript_status = st.empty()
                    transcript_status.info("Generating transcript with OpenAI Whisper...")
                    
                    transcript_result = transcribe_with_whisper(audio_result["file_path"], openai_api_key_input)
                    
                    if not transcript_result.get("success"):
                        st.error(f"Error generating transcript: {transcript_result.get('error', 'Unknown error')}")
                        continue
                    
                    transcript_status.success("✅ Transcript generated successfully")
                    
                    # Step 4: Fetch all comments
                    comment_status = st.empty()
                    comment_status.info("Fetching comments...")
                    
                    comments = get_all_comments_with_callback(video_id, youtube_api_key_input, status_text=comment_status)
                    
                    if not comments:
                        comment_status.warning("No comments retrieved or comments are disabled for this video.")
                        # Continue processing even if no comments are found
                    else:
                        comment_status.success(f"✅ Fetched {len(comments):,} comments")
                    
                    # Step 5: Create Word document with transcript and comments
                    doc_status = st.empty()
                    doc_status.info("Creating Word document...")
                    
                    doc_result = create_word_doc(
                        video_info, 
                        transcript_result["transcript"], 
                        comments
                    )
                    
                    if not doc_result.get("success"):
                        st.error(f"Error creating document: {doc_result.get('error', 'Unknown error')}")
                        continue
                    
                    doc_status.success("✅ Word document created successfully")
                    
                    # Create download link for the Word document
                    safe_title = clean_filename(video_info["title"])
                    doc_link = get_binary_file_downloader_html(
                        doc_result["file_path"], 
                        f"Download: {safe_title}_transcript_comments.docx"
                    )
                    
                    st.markdown(doc_link, unsafe_allow_html=True)
                    
                    # Store processed video in session state
                    st.session_state.processed_videos.append({
                        "video_id": video_id,
                        "title": video_info["title"],
                        "safe_title": safe_title,
                        "doc_path": doc_result["file_path"],
                        "url": url
                    })
                    
                    # Add a divider between videos
                    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
            
            # Show a message if videos were processed successfully
            if st.session_state.processed_videos:
                st.success(f"✅ Processed {len(st.session_state.processed_videos)} videos successfully! Go to the 'Analyze Videos' tab to analyze them.")

with tab2:
    # Function to handle the Analyze button click
    def on_analyze_click(video_idx):
        st.session_state.analyzing_video = video_idx
    
    # Display processed videos for analysis
    if not st.session_state.processed_videos:
        st.info("No videos have been processed yet. Please go to the 'Process Videos' tab to process videos first.")
    else:
        st.markdown("### Analyze Processed Videos")
        
        for idx, video_data in enumerate(st.session_state.processed_videos):
            with st.container():
                st.markdown(f"#### {idx+1}. {video_data['title']}")
                
                # Create doc link
                doc_link = get_binary_file_downloader_html(
                    video_data["doc_path"], 
                    f"Download Document"
                )
                st.markdown(doc_link, unsafe_allow_html=True)
                
                # Check if analysis already exists for this video
                video_id = video_data["video_id"]
                if video_id in st.session_state.analysis_results:
                    analysis_result = st.session_state.analysis_results[video_id]
                    
                    if "show" not in analysis_result:
                        analysis_result["show"] = True
                    
                    col1, col2 = st.columns([1, 3])
                    with col1:
                        if st.button("Toggle Analysis", key=f"toggle_{video_id}"):
                            analysis_result["show"] = not analysis_result["show"]
                    
                    with col2:
                        st.success("✅ Analysis completed")
                    
                    if analysis_result["show"]:
                        st.subheader("ChatGPT Analysis")
                        with st.expander("View Analysis", expanded=True):
                            st.markdown(analysis_result["analysis"])
                        
                        # Show download link if available
                        if "analysis_path" in analysis_result:
                            analysis_link = get_binary_file_downloader_html(
                                analysis_result["analysis_path"],
                                f"Download Analysis"
                            )
                            st.markdown(analysis_link, unsafe_allow_html=True)
                else:
                    # If currently analyzing this video
                    if st.session_state.analyzing_video == idx:
                        status_container = st.empty()
                        status_container.info("Analyzing video content with ChatGPT... This may take a minute.")
                        
                        # Perform the analysis
                        result = handle_analyze_document(
                            video_data["video_id"],
                            video_data["doc_path"],
                            video_data["safe_title"]
                        )
                        
                        # Store the result
                        if result["success"]:
                            st.session_state.analysis_results[video_id] = {
                                "success": True,
                                "analysis": result["analysis"],
                                "analysis_path": result["analysis_path"],
                                "analysis_filename": result["analysis_filename"],
                                "show": True
                            }
                            status_container.success("✅ Analysis completed successfully")
                            
                            # Show the analysis
                            st.subheader("ChatGPT Analysis")
                            with st.expander("View Analysis", expanded=True):
                                st.markdown(result["analysis"])
                            
                            # Show download link
                            analysis_link = get_binary_file_downloader_html(
                                result["analysis_path"],
                                f"Download Analysis"
                            )
                            st.markdown(analysis_link, unsafe_allow_html=True)
                        else:
                            st.session_state.analysis_results[video_id] = {
                                "success": False,
                                "error": result["error"]
                            }
                            status_container.error(f"❌ Error analyzing content: {result['error']}")
                        
                        # Reset analyzing state
                        st.session_state.analyzing_video = None
                    else:
                        # Show analyze button
                        if st.button("Analyze with ChatGPT", key=f"analyze_{video_id}", on_click=on_analyze_click, args=(idx,)):
                            # The actual analysis will happen on the next rerun, we just set the state here
                            pass
                
                # Add a divider between videos
                st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
